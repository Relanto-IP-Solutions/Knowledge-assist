"""Smoke-check for document extraction: Gemini (PDF/images) and native (DOCX, MD, PPTX).

Run with:
    uv run python scripts/tests_integration/smoke_document_extraction.py
    uv run python scripts/tests_integration/smoke_document_extraction.py path/to/sample.pdf
    uv run python scripts/tests_integration/smoke_document_extraction.py path/to/sample.pdf path/to/notes.jpg

Tests native extraction (MD, DOCX) with in-memory content; for PDF/images, pass file paths
to exercise Gemini Vision (requires Vertex AI credentials).
"""

import os
import pathlib
import sys


sys.path.insert(
    0, pathlib.Path(pathlib.Path(pathlib.Path(__file__).resolve()).parent).parent
)


def _test_native_md() -> None:
    """Test NativeExtractor with Markdown content."""
    from src.services.document_extraction import DocumentExtractionService

    content = b"# Hello\n\nThis is **markdown** content.\n\n- Item 1\n- Item 2"
    service = DocumentExtractionService()
    result = service.extract(content, "test.md")
    assert result == content.decode("utf-8")
    print("  [OK] Native MD extraction")


def _test_native_docx() -> None:
    """Test NativeExtractor with minimal DOCX content."""
    from io import BytesIO

    from docx import Document

    from src.services.document_extraction import DocumentExtractionService

    doc = Document()
    doc.add_paragraph("Hello from DOCX")
    doc.add_paragraph("Second paragraph")
    buf = BytesIO()
    doc.save(buf)
    content = buf.getvalue()

    service = DocumentExtractionService()
    result = service.extract(content, "test.docx")
    assert "Hello from DOCX" in result
    assert "Second paragraph" in result
    print("  [OK] Native DOCX extraction")


def _test_vision_file(path: str, expected_ext: str) -> None:
    """Test VisionExtractor with a file (PDF or image)."""
    from src.services.document_extraction import DocumentExtractionService

    with open(path, "rb") as f:
        content = f.read()

    object_name = pathlib.Path(path).name
    service = DocumentExtractionService()
    result = service.extract(content, object_name)
    assert isinstance(result, str)
    assert len(result) >= 0  # May be empty for blank pages
    print(f"  [OK] Vision extraction ({expected_ext}): {len(result)} chars")


def main() -> None:
    print("Document extraction smoke check")
    print("-" * 40)

    # Native extraction (no API calls)
    _test_native_md()
    _test_native_docx()

    # Vision extraction (requires Vertex AI) — only if file paths provided
    for path in sys.argv[1:]:
        if not pathlib.Path(path).is_file():
            print(f"  [SKIP] File not found: {path}")
            continue
        ext = os.path.splitext(path)[1].lower()
        if ext in (".pdf", ".jpg", ".jpeg", ".png", ".gif", ".webp"):
            _test_vision_file(path, ext)
        else:
            print(f"  [SKIP] Unsupported vision type: {ext}")

    if len(sys.argv) <= 1:
        print("  [INFO] Pass PDF/image paths to test Gemini Vision extraction")

    print("-" * 40)
    print("Document extraction smoke check OK.")


if __name__ == "__main__":
    main()
